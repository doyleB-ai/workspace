#!/usr/bin/env python3
"""
Generate a professional, publication-quality Word document from the
Comprehensive Financial Plan markdown file.
"""

import re
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------
NAVY = RGBColor(0x1B, 0x3A, 0x5C)
DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GRAY_HEX = "F2F2F2"
BORDER_GRAY_HEX = "CCCCCC"
CALLOUT_BG_HEX = "E8F0FE"
BLOCKQUOTE_BG_HEX = "F0F4FA"
NAVY_HEX = "1B3A5C"

BASE_DIR = Path(__file__).resolve().parent
MD_FILE = BASE_DIR / "COMPREHENSIVE-FINANCIAL-PLAN-2026-03-29.md"
CHARTS_DIR = BASE_DIR / "charts"
OUTPUT_FILE = BASE_DIR / "COMPREHENSIVE-FINANCIAL-PLAN-2026-03-29.docx"

FONT_BODY = "Calibri"
FONT_HEADING = "Calibri Light"


# ---------------------------------------------------------------------------
# Helper: set cell shading
# ---------------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    """Set thin borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:color'), color)
        element.set(qn('w:space'), '0')
        tcBorders.append(element)
    tcPr.append(tcBorders)


def set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for edge, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:w'), str(val))
        el.set(qn('w:type'), 'dxa')
        tcMar.append(el)
    tcPr.append(tcMar)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None, line_rule=None):
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    if before is not None:
        spacing.set(qn('w:before'), str(before))
    if after is not None:
        spacing.set(qn('w:after'), str(after))
    if line is not None:
        spacing.set(qn('w:line'), str(line))
    if line_rule is not None:
        spacing.set(qn('w:lineRule'), line_rule)


def add_bottom_border(paragraph, color=NAVY_HEX, size="6"):
    """Add a bottom border to a paragraph (for heading underlines)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:color'), color)
    bottom.set(qn('w:space'), '1')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_left_border(paragraph, color=NAVY_HEX, size="18"):
    """Add a left border to a paragraph (for blockquotes/callouts)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), size)
    left.set(qn('w:color'), color)
    left.set(qn('w:space'), '4')
    pBdr.append(left)


def set_paragraph_shading(paragraph, hex_color):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


def set_paragraph_indent(paragraph, left=None, first_line=None):
    pPr = paragraph._p.get_or_add_pPr()
    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    if left is not None:
        ind.set(qn('w:left'), str(left))
    if first_line is not None:
        ind.set(qn('w:firstLine'), str(first_line))


# ---------------------------------------------------------------------------
# Inline formatting: parse **bold** and *italic* from text
# ---------------------------------------------------------------------------
def add_formatted_runs(paragraph, text, font_name=FONT_BODY, font_size=Pt(11),
                       font_color=DARK_GRAY, base_bold=False, base_italic=False):
    """Parse markdown inline formatting and add runs to paragraph."""
    # Pattern to match **bold**, *italic*, ***bold italic***
    parts = re.split(r'(\*{1,3}[^*]+?\*{1,3})', text)
    for part in parts:
        if not part:
            continue
        m_bold_italic = re.match(r'^\*{3}(.+?)\*{3}$', part)
        m_bold = re.match(r'^\*{2}(.+?)\*{2}$', part)
        m_italic = re.match(r'^\*(.+?)\*$', part)

        if m_bold_italic:
            run = paragraph.add_run(m_bold_italic.group(1))
            run.bold = True
            run.italic = True
        elif m_bold:
            run = paragraph.add_run(m_bold.group(1))
            run.bold = True
            run.italic = base_italic
        elif m_italic:
            run = paragraph.add_run(m_italic.group(1))
            run.italic = True
            run.bold = base_bold
        else:
            run = paragraph.add_run(part)
            run.bold = base_bold
            run.italic = base_italic

        run.font.name = font_name
        run.font.size = font_size
        run.font.color.rgb = font_color
    return paragraph


def add_formatted_runs_to_cell(cell, text, font_name=FONT_BODY, font_size=Pt(10),
                               font_color=DARK_GRAY, bold=False, alignment=None):
    """Add formatted text to a table cell."""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    # Clear default
    for run in p.runs:
        run.text = ""
    add_formatted_runs(p, text, font_name=font_name, font_size=font_size,
                       font_color=font_color, base_bold=bold)
    set_paragraph_spacing(p, before=20, after=20, line=240, line_rule="auto")


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------
def parse_markdown(md_text):
    """Parse markdown into a list of elements."""
    elements = []
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # Skip the title line (# ...) - we handle it in cover page
        if i == 0 and line.startswith('# ') and not line.startswith('## '):
            i += 1
            continue

        # Skip metadata lines at top
        if i < 10 and (line.startswith('**Date:**') or line.startswith('**Prepared by:**') or
                       line.startswith('**Client:**')):
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            elements.append({'type': 'hr'})
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Headings
        h4_match = re.match(r'^####\s+(.*)', line)
        h3_match = re.match(r'^###\s+(.*)', line)
        h2_match = re.match(r'^##\s+(.*)', line)
        h1_match = re.match(r'^#\s+(.*)', line)

        if h4_match:
            elements.append({'type': 'h4', 'text': h4_match.group(1).strip()})
            i += 1
            continue
        elif h3_match:
            elements.append({'type': 'h3', 'text': h3_match.group(1).strip()})
            i += 1
            continue
        elif h2_match:
            elements.append({'type': 'h2', 'text': h2_match.group(1).strip()})
            i += 1
            continue
        elif h1_match:
            elements.append({'type': 'h1', 'text': h1_match.group(1).strip()})
            i += 1
            continue

        # Image
        img_match = re.match(r'^!\[([^\]]*)\]\(([^)]+)\)(\{[^}]*\})?', line)
        if img_match:
            alt = img_match.group(1)
            path = img_match.group(2)
            elements.append({'type': 'image', 'alt': alt, 'path': path})
            i += 1
            continue

        # Blockquote (may span multiple lines)
        if line.startswith('>'):
            bq_lines = []
            while i < len(lines) and lines[i].startswith('>'):
                bq_lines.append(lines[i].lstrip('>').strip())
                i += 1
            elements.append({'type': 'blockquote', 'text': ' '.join(bq_lines)})
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and re.match(r'^\|[-:|]+\|', lines[i + 1].strip()):
            table_lines = []
            while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            # Parse table
            headers = [c.strip() for c in table_lines[0].strip('|').split('|')]
            # Skip separator line
            rows = []
            for tl in table_lines[2:]:
                row = [c.strip() for c in tl.strip('|').split('|')]
                rows.append(row)
            elements.append({'type': 'table', 'headers': headers, 'rows': rows})
            continue

        # Bullet list
        list_match = re.match(r'^(\s*)([-*])\s+(.*)', line)
        if list_match:
            items = []
            while i < len(lines):
                lm = re.match(r'^(\s*)([-*])\s+(.*)', lines[i])
                if lm:
                    indent = len(lm.group(1))
                    level = indent // 2  # 0 for top-level, 1 for nested, etc.
                    items.append({'level': level, 'text': lm.group(3).strip()})
                    i += 1
                elif i < len(lines) and lines[i].strip() == '':
                    # Check if next non-empty line is also a list item
                    j = i + 1
                    while j < len(lines) and lines[j].strip() == '':
                        j += 1
                    if j < len(lines) and re.match(r'^(\s*)([-*])\s+', lines[j]):
                        i += 1  # skip empty line within list
                    else:
                        break
                else:
                    break
            elements.append({'type': 'list', 'items': items})
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', line)
        if num_match:
            items = []
            while i < len(lines):
                nm = re.match(r'^(\d+)\.\s+(.*)', lines[i])
                if nm:
                    items.append({'number': nm.group(1), 'text': nm.group(2).strip()})
                    i += 1
                elif lines[i].strip() == '':
                    j = i + 1
                    while j < len(lines) and lines[j].strip() == '':
                        j += 1
                    if j < len(lines) and re.match(r'^\d+\.\s+', lines[j]):
                        i += 1
                    else:
                        break
                else:
                    break
            elements.append({'type': 'numbered_list', 'items': items})
            continue

        # Regular paragraph
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() != '' and not lines[i].startswith('#') \
                and not lines[i].startswith('>') and not lines[i].startswith('|') \
                and not lines[i].startswith('---') and not lines[i].startswith('![') \
                and not re.match(r'^[-*]\s+', lines[i]) and not re.match(r'^\d+\.\s+', lines[i]):
            para_lines.append(lines[i])
            i += 1
        text = ' '.join(para_lines)
        elements.append({'type': 'paragraph', 'text': text})
        continue

    return elements


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------
def build_document(elements):
    doc = Document()

    # -----------------------------------------------------------------------
    # Page setup
    # -----------------------------------------------------------------------
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # -----------------------------------------------------------------------
    # Default font
    # -----------------------------------------------------------------------
    style = doc.styles['Normal']
    style.font.name = FONT_BODY
    style.font.size = Pt(11)
    style.font.color.rgb = DARK_GRAY
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = 1.15

    # -----------------------------------------------------------------------
    # Cover Page
    # -----------------------------------------------------------------------
    # Navy accent bar at top
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=0)
    # Create a thin navy line using a bordered paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top_border = OxmlElement('w:top')
    top_border.set(qn('w:val'), 'single')
    top_border.set(qn('w:sz'), '36')  # thick line
    top_border.set(qn('w:color'), NAVY_HEX)
    top_border.set(qn('w:space'), '1')
    pBdr.append(top_border)
    pPr.append(pBdr)

    # Spacer
    for _ in range(6):
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, before=0, after=0)

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title_p, before=0, after=120)
    run = title_p.add_run("Comprehensive Financial Plan")
    run.font.name = FONT_HEADING
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = False

    # Subtitle
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(sub_p, before=0, after=80)
    run = sub_p.add_run("Cleary Family")
    run.font.name = FONT_HEADING
    run.font.size = Pt(18)
    run.font.color.rgb = NAVY

    # Thin navy line
    line_p = doc.add_paragraph()
    set_paragraph_spacing(line_p, before=0, after=200)
    add_bottom_border(line_p, color=NAVY_HEX, size="6")

    # Date
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(date_p, before=0, after=60)
    run = date_p.add_run("March 29, 2026")
    run.font.name = FONT_BODY
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_GRAY

    # Prepared by
    prep_p = doc.add_paragraph()
    prep_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(prep_p, before=0, after=60)
    run = prep_p.add_run("Prepared by: Fiduciary Advisory Board")
    run.font.name = FONT_BODY
    run.font.size = Pt(13)
    run.font.color.rgb = DARK_GRAY

    # Spacers to push confidential to bottom
    for _ in range(10):
        sp = doc.add_paragraph()
        set_paragraph_spacing(sp, before=0, after=0)

    # Confidential notice
    conf_p = doc.add_paragraph()
    conf_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(conf_p, before=0, after=0)
    run = conf_p.add_run("CONFIDENTIAL")
    run.font.name = FONT_BODY
    run.font.size = Pt(9)
    run.font.color.rgb = DARK_GRAY
    run.italic = True
    run2 = conf_p.add_run(
        " — This document contains sensitive financial information prepared "
        "exclusively for the Cleary family. It is not intended for distribution, "
        "reproduction, or use by any third party."
    )
    run2.font.name = FONT_BODY
    run2.font.size = Pt(9)
    run2.font.color.rgb = DARK_GRAY
    run2.italic = True

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # Set up header and footer (applies to all pages after cover)
    # -----------------------------------------------------------------------
    # We need a separate section for content so cover has no header/footer
    new_section = doc.add_section()
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)
    new_section.top_margin = Inches(1)
    new_section.bottom_margin = Inches(1)
    new_section.left_margin = Inches(1)
    new_section.right_margin = Inches(1)

    # Unlink from previous section
    new_section.header.is_linked_to_previous = False
    new_section.footer.is_linked_to_previous = False

    # First section: no header/footer
    first_section = doc.sections[0]
    first_section.different_first_page_header_footer = False
    # Clear first section header/footer
    for p in first_section.header.paragraphs:
        p.text = ""
    for p in first_section.footer.paragraphs:
        p.text = ""

    # Second section header
    header = new_section.header
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Cleary Family — Comprehensive Financial Plan")
    run.font.name = FONT_BODY
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_GRAY
    run.italic = True

    # Second section footer
    footer = new_section.footer
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    # CONFIDENTIAL left-aligned
    run_conf = fp.add_run("CONFIDENTIAL")
    run_conf.font.name = FONT_BODY
    run_conf.font.size = Pt(7)
    run_conf.font.color.rgb = DARK_GRAY
    run_conf.font.small_caps = True

    # Page number - add tab and page field
    run_tab = fp.add_run("\t\t")
    run_tab.font.size = Pt(8)

    # Add page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_page = fp.add_run()
    run_page.font.name = FONT_BODY
    run_page.font.size = Pt(8)
    run_page.font.color.rgb = DARK_GRAY
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)

    # Set footer tab stops for center and right alignment
    pPr = fp._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    center_tab = OxmlElement('w:tab')
    center_tab.set(qn('w:val'), 'center')
    center_tab.set(qn('w:pos'), '4680')  # center of 6.5" content width
    tabs.append(center_tab)
    right_tab = OxmlElement('w:tab')
    right_tab.set(qn('w:val'), 'right')
    right_tab.set(qn('w:pos'), '9360')
    tabs.append(right_tab)
    pPr.append(tabs)

    # -----------------------------------------------------------------------
    # Render elements
    # -----------------------------------------------------------------------
    skip_initial_confidential = True  # Skip the first blockquote (confidential notice)

    for idx, elem in enumerate(elements):
        etype = elem['type']

        if etype == 'hr':
            # Thin horizontal rule
            hr_p = doc.add_paragraph()
            set_paragraph_spacing(hr_p, before=120, after=120)
            add_bottom_border(hr_p, color=BORDER_GRAY_HEX, size="4")
            continue

        if etype == 'h2':
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=480, after=240)
            add_bottom_border(p, color=NAVY_HEX, size="6")
            text = elem['text']
            # Remove markdown formatting for heading
            clean_text = re.sub(r'\*{1,3}', '', text)
            run = p.add_run(clean_text)
            run.font.name = FONT_HEADING
            run.font.size = Pt(18)
            run.font.color.rgb = NAVY
            run.bold = True
            continue

        if etype == 'h3':
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=360, after=160)
            text = elem['text']
            clean_text = re.sub(r'\*{1,3}', '', text)
            run = p.add_run(clean_text)
            run.font.name = FONT_HEADING
            run.font.size = Pt(14)
            run.font.color.rgb = NAVY
            run.bold = True
            continue

        if etype == 'h4':
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=240, after=120)
            text = elem['text']
            clean_text = re.sub(r'\*{1,3}', '', text)
            run = p.add_run(clean_text)
            run.font.name = FONT_BODY
            run.font.size = Pt(12)
            run.font.color.rgb = DARK_GRAY
            run.bold = True
            run.italic = True
            continue

        if etype == 'paragraph':
            text = elem['text']
            # Check if this is a callout (Bottom Line or Our Recommendation)
            is_callout = False
            callout_patterns = ['**Bottom Line:**', '**Our Recommendation:**']
            for cp in callout_patterns:
                if cp in text:
                    is_callout = True
                    break

            if is_callout:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=200, after=200, line=276, line_rule="auto")
                set_paragraph_shading(p, CALLOUT_BG_HEX)
                add_left_border(p, color=NAVY_HEX, size="18")
                set_paragraph_indent(p, left=360)
                add_formatted_runs(p, text, font_size=Pt(11), font_color=DARK_GRAY)
            else:
                p = doc.add_paragraph()
                set_paragraph_spacing(p, before=0, after=120, line=276, line_rule="auto")
                add_formatted_runs(p, text)
            continue

        if etype == 'blockquote':
            text = elem['text']
            if skip_initial_confidential and 'CONFIDENTIAL' in text:
                skip_initial_confidential = False
                continue

            # Check if this is a callout
            is_callout = False
            callout_patterns = ['**Bottom Line:**', '**Our Recommendation:**']
            for cp in callout_patterns:
                if cp in text:
                    is_callout = True
                    break

            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=200, after=200, line=276, line_rule="auto")
            if is_callout:
                set_paragraph_shading(p, CALLOUT_BG_HEX)
            else:
                set_paragraph_shading(p, BLOCKQUOTE_BG_HEX)
            add_left_border(p, color=NAVY_HEX, size="18")
            set_paragraph_indent(p, left=360)
            add_formatted_runs(p, text, font_size=Pt(11), font_color=DARK_GRAY,
                               base_italic=not is_callout)
            continue

        if etype == 'image':
            img_path = CHARTS_DIR / os.path.basename(elem['path'])
            if not img_path.exists():
                # Try relative path from markdown
                img_path = BASE_DIR / elem['path']
            if img_path.exists():
                # Center the image
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, before=200, after=80)
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6))

                # Caption
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(cap_p, before=0, after=200)
                alt_text = elem.get('alt', 'Chart')
                run = cap_p.add_run(alt_text)
                run.font.name = FONT_BODY
                run.font.size = Pt(9)
                run.font.color.rgb = DARK_GRAY
                run.italic = True
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"[Chart not found: {elem['path']}]")
                run.font.name = FONT_BODY
                run.font.size = Pt(9)
                run.font.color.rgb = DARK_GRAY
                run.italic = True
            continue

        if etype == 'table':
            headers = elem['headers']
            rows = elem['rows']
            num_cols = len(headers)

            table = doc.add_table(rows=1 + len(rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = True

            # Set table width to full page
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '5000')
            tblW.set(qn('w:type'), 'pct')
            tblPr.append(tblW)

            # Header row
            for j, header_text in enumerate(headers):
                cell = table.rows[0].cells[j]
                set_cell_shading(cell, NAVY_HEX)
                set_cell_borders(cell, color=NAVY_HEX)
                set_cell_margins(cell)
                clean = re.sub(r'\*{1,3}', '', header_text)
                add_formatted_runs_to_cell(cell, clean, font_size=Pt(10),
                                           font_color=WHITE, bold=True)

            # Data rows
            for ri, row in enumerate(rows):
                for j in range(min(len(row), num_cols)):
                    cell = table.rows[ri + 1].cells[j]
                    # Alternating row colors
                    if ri % 2 == 1:
                        set_cell_shading(cell, LIGHT_GRAY_HEX)
                    set_cell_borders(cell, color=BORDER_GRAY_HEX)
                    set_cell_margins(cell)

                    cell_text = row[j]
                    clean = cell_text.strip()

                    # Determine alignment - numbers/currency right-aligned
                    is_number = bool(re.match(r'^[\$\-~+]?[\d,]+\.?\d*[%MKBmkb]?$',
                                              clean.replace('**', '').replace('*', '').replace('~', '').strip()))
                    is_bold = clean.startswith('**') and clean.endswith('**')

                    align = WD_ALIGN_PARAGRAPH.RIGHT if is_number else WD_ALIGN_PARAGRAPH.LEFT
                    add_formatted_runs_to_cell(cell, clean, font_size=Pt(10),
                                               font_color=DARK_GRAY, bold=is_bold,
                                               alignment=align)

            # Add spacing after table
            spacer = doc.add_paragraph()
            set_paragraph_spacing(spacer, before=0, after=120)
            continue

        if etype == 'list':
            for item in elem['items']:
                p = doc.add_paragraph()
                level = item['level']
                indent = 360 + (level * 360)  # 0.25 inch per level in twips
                set_paragraph_indent(p, left=indent)
                set_paragraph_spacing(p, before=0, after=60, line=276, line_rule="auto")

                # Bullet character
                bullet = "•  " if level == 0 else "–  "
                run = p.add_run(bullet)
                run.font.name = FONT_BODY
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_GRAY

                add_formatted_runs(p, item['text'])
            continue

        if etype == 'numbered_list':
            for item in elem['items']:
                p = doc.add_paragraph()
                set_paragraph_indent(p, left=360)
                set_paragraph_spacing(p, before=0, after=80, line=276, line_rule="auto")

                run = p.add_run(f"{item['number']}.  ")
                run.font.name = FONT_BODY
                run.font.size = Pt(11)
                run.font.color.rgb = DARK_GRAY
                run.bold = True

                add_formatted_runs(p, item['text'])
            continue

    return doc


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main():
    print(f"Reading markdown from: {MD_FILE}")
    md_text = MD_FILE.read_text(encoding='utf-8')

    print("Parsing markdown...")
    elements = parse_markdown(md_text)
    print(f"  Found {len(elements)} elements")

    print("Building document...")
    doc = build_document(elements)

    print(f"Saving to: {OUTPUT_FILE}")
    doc.save(str(OUTPUT_FILE))

    size = OUTPUT_FILE.stat().st_size
    print(f"Done! File size: {size:,} bytes ({size/1024:.1f} KB)")


if __name__ == '__main__':
    main()
